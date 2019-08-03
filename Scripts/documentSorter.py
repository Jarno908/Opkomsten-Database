#! Python3
import logging
logging.basicConfig(level=logging.DEBUG, format=' %(asctime)s - %(levelname)s - %(message)s')
log = logging.getLogger(__name__)

from pathlib2 import Path
import docx
import constants
import documentClass
from opkomst_module import Opkomst

#from pprint import pprint

config = None

def setup_documentSorter(config_dict):
    global config
    config = config_dict

# Sorteert de bestanden in de input folder en maakt een file_path voor ieder
def sort(files_list):

    sorted_documents = []

    for file in files_list:
        if file.suffix == ".docx":
            log.info("Currently being sorted: {}".format(file))
            document_data = getDocumentData(str(file))
            #pprint(document_data)

            if isinstance(document_data, documentClass.Document) and document_data.template_version != 0:
                log.info("Document uses {}{} format".format(document_data.document_type, document_data.template_version))

                if document_data.document_type == "Opkomst":
                    custom_path = opkomstPath(document_data)
                else:
                    raise Exception("Not a supported document-type!")

                if document_data.titel != "":
                    custom_path = custom_path.joinpath(document_data.titel)
                else:
                    custom_path = custom_path.joinpath(file.stem)

                custom_path = custom_path.with_suffix(".docx")
                document_data.file_path = str(custom_path)
                sorted_documents.append(document_data)

            else:
                log.debug("File {} does not use a valid template".format(file.name))

    return sorted_documents

# Alle informatie wordt uit het document gehaald en in een Document object gesptopt
def getDocumentData(file_path):
    try:
        doc = docx.Document(file_path)
    except ValueError:
        log.error('File provided is not a docx file!')
        return None
    except:
        log.error('Read Docx Error!!!')
        return None
    else:
        global config
        data = {}
        data["local_path"] = file_path
        if  (len(doc.sections[0].header.tables) > 0
            and doc.sections[0].header.tables[0].cell(0, 0).paragraphs[0].text == "Template-versie"):
                data["version"] = doc.sections[0].header.tables[0].cell(0, 1).paragraphs[0].text
        else:
            data["version"] = doc.core_properties.version

        if len(data["version"]) != 4:
            return None
        else:
            table = doc.tables[0]
            keys = []

            for cell in table.column_cells(0):
                keys.append(cell.paragraphs[0].text)

            for i in range(len(keys)):
                fullText = []
                for paragraph in table.column_cells(1)[i].paragraphs:
                    if keys[i] == "Omschrijving":
                        fullText.append(paragraph.text)
                    else:
                        fullText.append(paragraph.text.replace("\n", "").replace(",", "\n"))

                fullText = "\n".join(fullText)

                data[keys[i]] = fullText

            if "Titel" in data:
                for c in constants.FORBIDDEN_CHARACTERS:
                    data["Titel"] = data["Titel"].replace(c, "")

            matching_keys = data.keys() & constants.KEYS_TO_LOWER
            for key in matching_keys:
                data[key] = data[key].lower()

            if config != None:
                data["Uploader_Name"] = config["Preferences"]["uploader_name"]

            data["document_type"] = constants.DOCUMENT_TYPES.get(data["version"][0])
            data["version"] = int(data["version"][1:])

            if data["document_type"] == "Opkomst":
                return Opkomst(data)

            return None

def opkomstPath(document_data):
    custom_path = Path("/Documenten Reggegroep/Opkomsten")

    if document_data.speltakken[0] in constants.SPELTAKKEN:
        custom_path = custom_path.joinpath(document_data.speltakken[0])
    else:
        custom_path = custom_path.joinpath("Overig")

    if document_data.categorie in constants.CATEGORIES:
        custom_path = custom_path.joinpath(document_data.categorie)
    else:
        custom_path = custom_path.joinpath("Overig")

    return custom_path
